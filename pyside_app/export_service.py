from docx import Document
from docx.shared import Pt
import os
import markdown

class ExportService:
    def __init__(self, template_path="resources/templates/default.docx"):
        self.template_path = template_path

    def render_markdown_to_docx(self, md_content: str, output_path: str):
        """
        Takes raw markdown, parses it, and injects it into a Word template.
        For simplicity in this demonstration, this generates a basic docx,
        using the template if it exists.
        """
        # Load template if it explicitly exists as a file, and catch potential format errors
        doc = None
        if self.template_path:
            if os.path.isfile(self.template_path):
                try:
                    doc = Document(self.template_path)
                except Exception as e:
                    # If explicit template was passed but failed to load (e.g., bad format or not a docx),
                    # raise an error instead of silently falling back to a blank document.
                    raise ValueError(f"Failed to load specified template {self.template_path}: {e}") from e
            else:
                raise ValueError(f"Specified template file does not exist: {self.template_path}")
        else:
            doc = Document()

        # Parse markdown into HTML string
        html_str = markdown.markdown(md_content)

        # A full production implementation would parse HTML AST to docx runs here.
        # Here we perform basic parsing of markdown headings for demonstration.
        for line in md_content.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('# '):
                heading = doc.add_heading(level=1)
                run = heading.add_run(line[2:])
                run.font.size = Pt(16)
            elif line.startswith('## '):
                heading = doc.add_heading(level=2)
                run = heading.add_run(line[3:])
                run.font.size = Pt(14)
            elif line.startswith('### '):
                heading = doc.add_heading(level=3)
                run = heading.add_run(line[4:])
                run.font.size = Pt(12)
            else:
                doc.add_paragraph(line)

        doc.save(output_path)
        return output_path
